"""
CORNELSEN EBOOK - NUR BUCHINHALT → WORD
========================================

Dieses Skript macht Screenshots NUR vom Buchinhalt-Bereich
(die zwei Seiten in der Mitte - ohne Header/Navigation)
und speichert sie in einer Word-Datei (Querformat)

VERWENDUNG:
1. Login-Daten sind bereits eingetragen (lisibrinki)
2. Führe aus: python cornelsen_buchinhalt.py
3. Gib den gewünschten Dateinamen ein (z.B. "MeinEbook")
4. Browser öffnet sich → Login läuft
5. ⏸️  PROGRAMM WARTET AUF DICH!
6. Navigiere zur Startseite im Browser
7. Drücke ENTER im Terminal
8. Screenshots werden erstellt
9. Word-Datei wird im aktuellen Verzeichnis gespeichert
10. Fertig!

FEATURES:
✅ Interaktive Dateinamen-Eingabe
✅ WARTET auf dein Signal! ⭐
✅ Querformat (Landscape)
✅ Nur Buchinhalt (ohne Header/Navigation)
✅ Login bereits konfiguriert
✅ Speichert im aktuellen Verzeichnis
"""

from playwright.sync_api import sync_playwright
from docx import Document
from docx.shared import Cm
from docx.enum.section import WD_ORIENT
import time
import os

# ============================================================================
# KONFIGURATION
# ============================================================================

#EMAIL = "lisibrinki"
#PASSWORT = "Buchholz25-"
#EBOOK_URL = "https://mein.westermann.de/auth/login?scope=openid&state=0b8945846a8b4dc22d94976800949df3&response_type=code&approval_prompt=auto&redirect_uri=https://www.westermann.de/backend/oauth2/post-login&client_id=MH8rAJnf"

#Buch 3
#EMAIL = "lisa.baltzer@bbs-buchholz.de"
#PASSWORT = "Buchholz25!-"
#EBOOK_URL = "https://www.westermann.de/meine-produkte"
#EBOOK_URL = "https://bibox2.westermann.de/v2/book/7365/page/1"

EMAIL = "lisibrinki"
PASSWORT = "Buchholz26-!"
EBOOK_URL = "https://ebook.cornelsen.de/auth/220053709/ebook"


#Startseite und Anzahl der Seiten
STARTSEITE = 1
ANZAHL_SEITEN = 300  # Anzahl der Seiten, die du speichern möchtest

# Dateiname wird beim Start abgefragt!
ORDNER_SCREENSHOTS = "screenshots_buchinhalt"

WARTEZEIT_ZWISCHEN_SEITEN = 2
BROWSER_SICHTBAR = True

# BROWSERGRÖSSE (bei Bedarf anpassen)
BROWSER_WIDTH = 1920    # Browserbreite (verkleinern, wenn zu groß)
BROWSER_HEIGHT = 1080   # Browserhöhe (verkleinern, wenn zu groß)
# Für kleinere Bildschirme: Versuchen Sie 1366x768 oder 1280x720

# ============================================================================
# BUCHINHALT-SELEKTOREN (Versuche verschiedene Möglichkeiten)
# ============================================================================

# Diese Selektoren werden in der Reihenfolge ausprobiert
# Der erste funktionierende wird verwendet

BUCHINHALT_SELEKTOREN = [
    # Methode 1: Direkt das Seiten-Container-Element
    '.page-container',
    '.pages-container', 
    '.book-pages',
    '.reader-pages',
    
    # Methode 2: iFrame mit Buchinhalt
    'iframe',
    
    # Methode 3: Main-Container
    'main',
    '[role="main"]',
    '#main-content',
    
    # Methode 4: Spezifische Klassen
    '.ebook-reader',
    '.reader-content',
    '#reader',
    '#book-content',
    
    # Methode 5: Canvas (falls PDF-basiert)
    'canvas',
]

# ============================================================================

def detecter_zone_automatique(page):
    """
    Erkennt automatisch den Seitenbereich im Browser
    Funktioniert unabhängig von der Bildschirmgröße
    """
    try:
        result = page.evaluate("""
            () => {
                const selectors = [
                    'canvas', '.page', '[class*="page"]', 'img[src*="page"]',
                    'iframe', '[role="main"]', 'main', '.reader', '.book',
                    '[class*="viewer"]', '[class*="reader"]', '[class*="book"]'
                ];

                let bestElement = null;
                let maxArea = 0;

                for (const selector of selectors) {
                    try {
                        const elements = document.querySelectorAll(selector);
                        for (const el of elements) {
                            const rect = el.getBoundingClientRect();
                            const area = rect.width * rect.height;

                            if (area > 100000 && rect.top >= 0 && rect.left >= 0 &&
                                rect.width > 500 && rect.height > 500) {
                                if (area > maxArea) {
                                    maxArea = area;
                                    bestElement = {
                                        x: Math.round(rect.left),
                                        y: Math.round(rect.top),
                                        width: Math.round(rect.width),
                                        height: Math.round(rect.height)
                                    };
                                }
                            }
                        }
                    } catch (e) {}
                }

                // Fallback: zone centrale
                if (!bestElement) {
                    const w = window.innerWidth;
                    const h = window.innerHeight;
                    bestElement = {
                        x: Math.round(w * 0.1),
                        y: Math.round(h * 0.1),
                        width: Math.round(w * 0.8),
                        height: Math.round(h * 0.8)
                    };
                }

                return bestElement;
            }
        """)
        return result
    except:
        return None


def finde_buchinhalt_element(page):
    """Findet das Element mit dem Buchinhalt"""

    print("🔍 Suche Buchinhalt-Element...")

    for selektor in BUCHINHALT_SELEKTOREN:
        try:
            element = page.locator(selektor).first

            if element.count() > 0:
                # Prüfe ob Element sichtbar ist
                if element.is_visible():
                    print(f"   ✓ Gefunden mit Selektor: {selektor}")
                    return element, selektor
        except:
            continue

    print("   ⚠️ Kein spezifisches Element gefunden")
    print("   → Nutze Fallback: Détection automatique")
    return None, None


def screenshot_mit_koordinaten(page, pfad):
    """
    Screenshot von einem bestimmten Bereich mit fixen Koordinaten
    """

    # FIXE KOORDINATEN - Nur die 2 Buchseiten ohne Ränder
    clip = {
        'x': 230,      # Abstand vom linken Rand > nach links schieben
        'y': 20,       # Abstand von oben > nach oben schieben
        'width': 1455, # Breite der Erfassungszone > breiter machen
        'height': 1000  # Höhe der Erfassungszone
    }

    print(f"   ✓ Nutze fixe Koordinaten: {clip['width']}×{clip['height']}px")

    page.screenshot(path=pfad, clip=clip)
    return clip


def warte_auf_seiten_stabilitaet(page, wartezeit=3.0):
    """
    Wartet bis die Seite stabil und scharf ist
    Hilft gegen verschwommene Screenshots
    """
    print("   → Warte auf Seiten-Stabilität...", end="", flush=True)

    # Warte die Mindestzeit
    time.sleep(wartezeit)

    # Zusätzliche Prüfung: Warte bis keine Netzwerkaktivität mehr
    try:
        page.wait_for_load_state("networkidle", timeout=3000)
    except:
        pass  # Timeout ist ok, wir haben schon gewartet

    print(" ✓")


def hauptprogramm():
    """Hauptprogramm - Screenshots nur vom Buchinhalt"""
    
    print("\n" + "="*70)
    print("  CORNELSEN EBOOK - NUR BUCHINHALT → WORD")
    print("="*70 + "\n")
    
    if EMAIL == "deine.email@example.com":
        print("❌ FEHLER: Bitte trage zuerst deine Daten ein!")
        return
    
    # ===== DATEINAME ABFRAGEN =====
    print("📝 Word-Dateiname:")
    print("="*70 + "\n")
    
    dateiname_eingabe = input("Dateiname (ohne .docx): ").strip()
    
    # Falls leer, nutze Standard
    if not dateiname_eingabe:
        dateiname_eingabe = "eBook"
        print(f"   → Nutze Standard: {dateiname_eingabe}")
    
    # .docx hinzufügen falls nicht vorhanden
    if not dateiname_eingabe.endswith('.docx'):
        WORD_DATEINAME = f"{dateiname_eingabe}.docx"
    else:
        WORD_DATEINAME = dateiname_eingabe
    
    print(f"   ✓ Dateiname: {WORD_DATEINAME}\n")
    
    print("📋 Konfiguration:")
    print(f"   Email: {EMAIL}")
    print(f"   Seiten: {STARTSEITE} bis {STARTSEITE + ANZAHL_SEITEN - 1}")
    print(f"   Ausgabe: {WORD_DATEINAME} (Querformat)\n")
    
    os.makedirs(ORDNER_SCREENSHOTS, exist_ok=True)
    screenshots = []
    
    with sync_playwright() as p:
        print("🚀 Starte Browser...")
        browser = p.chromium.launch(
            headless=not BROWSER_SICHTBAR,
            args=[
                '--force-device-scale-factor=1.0',  # Force zoom à 100%
                '--disable-blink-features=AutomationControlled'
            ]
        )

        # Créer un contexte avec zoom fixe à 100%
        context = browser.new_context(
            viewport={'width': BROWSER_WIDTH, 'height': BROWSER_HEIGHT},
            device_scale_factor=1.0,  # Force l'échelle à 1.0 (100%)
            no_viewport=False
        )
        page = context.new_page()

        print(f"   → Fenêtre: {BROWSER_WIDTH}×{BROWSER_HEIGHT} @ 100% zoom")
        
        try:
            # ===== LOGIN =====
            print("🔐 Anmeldung...")
            page.goto(EBOOK_URL, wait_until="networkidle")
            time.sleep(2)

            try:
                page.fill('input[type="email"]', EMAIL)
                page.fill('input[type="password"]', PASSWORT)
                page.click('button[type="submit"]')
                time.sleep(3)
                print("   ✓ Angemeldet\n")
            except:
                print("   ⚠️ Manuelle Anmeldung erforderlich\n")
            
            # ===== WARTE AUF BENUTZER-SIGNAL =====
            print("="*70)
            print("⏸️  WARTE AUF DEIN SIGNAL!")
            print("="*70 + "\n")
            
            print("Der Browser ist jetzt offen.")
            print("\n📋 JETZT BIST DU DRAN:")
            print("   1. Prüfe ob du eingeloggt bist")
            print("   2. Falls nicht: Melde dich manuell an")
            print("   3. Navigiere zur STARTSEITE im Browser")
            print(f"      (Seite {STARTSEITE} deines eBooks)")
            print("   4. Stelle sicher, dass die Buchseiten sichtbar sind")
            print("   5. Optional: Aktiviere Vollbild (F11) für bessere Qualität")
            print("   6. Wenn alles bereit ist: Drücke ENTER hier im Terminal\n")
            
            print("💡 TIPP:")
            print("   Im Vollbild (F11) werden die Buchseiten größer")
            print("   → Bessere Screenshot-Qualität!\n")
            
            print("="*70)
            input("➡️  DRÜCKE ENTER WENN DU BEREIT BIST...")
            print("="*70 + "\n")
            
            # Warte auf eBook-Inhalt
            print("   → Warte kurz auf eBook-Inhalt...")
            time.sleep(2)

            # ===== AJUSTER AUTOMATIQUEMENT LE ZOOM =====
            print("   → Ajustement automatique du zoom...")
            try:
                # Réinitialiser le zoom du navigateur à 100%
                page.evaluate("document.body.style.zoom = '1.0'")

                # Forcer le zoom à 100% avec Ctrl+0
                page.keyboard.press("Control+0")
                time.sleep(0.5)

                print("   ✓ Zoom ajusté à 100%")
            except Exception as e:
                print(f"   ⚠️ Zoom manuel requis (Ctrl+0)")

            print("   ✓ Bereit!\n")

            # ===== VERWENDE FIXE KOORDINATEN =====
            print("✅ Verwende fixe Koordinaten-Screenshot")
            print(f"   → Koordinaten: x=290, y=50, width=1340, height=920\n")
            screenshot_methode = "koordinaten"
            buchinhalt_element = None
            verwendeter_selektor = None
            
            # ===== SCREENSHOTS ERSTELLEN =====
            print(f"📸 Erstelle Screenshots (Methode: {screenshot_methode})...")
            print("="*70)
            
            erfolgreiche_screenshots = 0
            
            for i in range(ANZAHL_SEITEN):
                aktuelle_seite = STARTSEITE + i

                # Fortschritt
                prozent = (i + 1) / ANZAHL_SEITEN * 100
                balken = "█" * int(prozent / 5) + "░" * (20 - int(prozent / 5))
                print(f"[{balken}] {prozent:5.1f}% | Seite {aktuelle_seite:3d}")

                try:
                    # Warte bis die Seite stabil und scharf ist
                    # Besonders wichtig für erste Seite die oft verschwommen startet
                    if i == 0:
                        # Erste Seite braucht mehr Zeit (oft verschwommen)
                        warte_auf_seiten_stabilitaet(page, wartezeit=4.0)
                    else:
                        # Nachfolgende Seiten sind schneller stabil
                        warte_auf_seiten_stabilitaet(page, wartezeit=2.5)

                    screenshot_pfad = f"{ORDNER_SCREENSHOTS}/seite_{aktuelle_seite:04d}.png"

                    # Screenshot erstellen
                    if screenshot_methode == "element" and buchinhalt_element:
                        # Element-Screenshot (nur Buchinhalt)
                        buchinhalt_element.screenshot(path=screenshot_pfad, timeout=10000)
                    else:
                        # Koordinaten-Screenshot
                        screenshot_mit_koordinaten(page, screenshot_pfad)

                    screenshots.append(screenshot_pfad)
                    erfolgreiche_screenshots += 1
                    print("   ✓ Screenshot erfolgreich")

                except Exception as e:
                    print(f"   ✗ Fehler: {str(e)[:50]}")

                    # Fallback: Vollbild-Screenshot
                    try:
                        screenshot_pfad = f"{ORDNER_SCREENSHOTS}/seite_{aktuelle_seite:04d}.png"
                        page.screenshot(path=screenshot_pfad)
                        screenshots.append(screenshot_pfad)
                        print(f"   → Fallback: Vollbild-Screenshot")
                    except:
                        pass

                # Zur nächsten Seite navigieren (einfache Pfeiltaste)
                if i < ANZAHL_SEITEN - 1:
                    print(f"   → Navigation zu Seite {aktuelle_seite + 1}...")
                    try:
                        page.keyboard.press("ArrowRight")
                        time.sleep(WARTEZEIT_ZWISCHEN_SEITEN)
                        print("   ✓ Navigation erfolgreich")
                    except Exception as e:
                        print(f"   ⚠️ Navigation fehlgeschlagen: {str(e)[:50]}")
                    
                    print("")  # Leerzeile für bessere Lesbarkeit
            
            print(f"\n{'='*70}")
            print(f"✓ {erfolgreiche_screenshots} von {ANZAHL_SEITEN} Screenshots erstellt!\n")
            
        except Exception as e:
            print(f"\n❌ Fehler: {e}")
            import traceback
            traceback.print_exc()
            return
        
        finally:
            browser.close()

    # ===== WORD DOKUMENT ERSTELLEN =====
    if not screenshots:
        print("❌ Keine Screenshots vorhanden")
        return
    
    print("📝 Erstelle Word-Dokument (Querformat)...")
    print("="*70)
    
    try:
        doc = Document()
        
        # Querformat einstellen
        section = doc.sections[0]
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7)   # A4 Querformat
        section.page_height = Cm(21.0)
        
        # Minimale Ränder
        section.top_margin = Cm(0.5)
        section.bottom_margin = Cm(0.5)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)
        
        # Screenshots einfügen - OHNE leere Seiten
        for i, screenshot in enumerate(screenshots, 1):
            prozent = i / len(screenshots) * 100
            balken = "█" * int(prozent / 5) + "░" * (20 - int(prozent / 5))
            print(f"[{balken}] {prozent:5.1f}% | Bild {i}/{len(screenshots)}", end="\r")

            try:
                # Ersten Screenshot direkt einfügen, danach neue Sektion für jedes Bild
                if i == 1:
                    # Erstes Bild: Direkt in erste Seite einfügen
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(screenshot, width=Cm(28.7))

                    # Keine Abstände
                    paragraph.paragraph_format.space_before = Cm(0)
                    paragraph.paragraph_format.space_after = Cm(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                else:
                    # Weitere Bilder: Neue Sektion erstellen (verhindert leere Seiten)
                    new_section = doc.add_section(WD_ORIENT.LANDSCAPE)
                    new_section.page_width = Cm(29.7)
                    new_section.page_height = Cm(21.0)
                    new_section.top_margin = Cm(0.5)
                    new_section.bottom_margin = Cm(0.5)
                    new_section.left_margin = Cm(0.5)
                    new_section.right_margin = Cm(0.5)

                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()
                    run.add_picture(screenshot, width=Cm(28.7))

                    # Keine Abstände
                    paragraph.paragraph_format.space_before = Cm(0)
                    paragraph.paragraph_format.space_after = Cm(0)
                    paragraph.paragraph_format.line_spacing = 1.0

            except Exception as e:
                print(f"\n   ⚠️ Fehler bei Bild {i}: {e}")
        
        # Word-Dokument speichern
        doc.save(WORD_DATEINAME)
        dateigröße = os.path.getsize(WORD_DATEINAME) / (1024 * 1024)
        
        # Vollständiger Pfad
        vollständiger_pfad = os.path.abspath(WORD_DATEINAME)
        
        print(f"\n{'='*70}")
        print(f"✓ Word-Dokument erstellt: {WORD_DATEINAME}")
        print(f"  📁 Speicherort: {vollständiger_pfad}")
        print(f"  💾 Größe: {dateigröße:.1f} MB")
        print(f"  📄 Seiten: {len(screenshots)}")
        print(f"  📐 Format: Querformat (Landscape)")
        
    except Exception as e:
        print(f"\n❌ Fehler beim Word-Erstellen: {e}")
        import traceback
        traceback.print_exc()
        return
    
    print("\n" + "="*70)
    print("✅ FERTIG!")
    print("="*70)
    print(f"\nDein eBook: {WORD_DATEINAME}")
    print(f"Speicherort: {os.path.abspath(WORD_DATEINAME)}")
    print(f"Screenshots: {ORDNER_SCREENSHOTS}/\n")
    
    print("💡 DATEI ÖFFNEN:")
    print(f"   → Im Explorer: explorer {os.path.dirname(os.path.abspath(WORD_DATEINAME)) or '.'}")
    print(f"   → Direkt öffnen: start {WORD_DATEINAME}\n")
    
    print("💡 WICHTIG:")
    print("   Schaue dir die erste Screenshot-Datei an:")
    print(f"   {ORDNER_SCREENSHOTS}/seite_{STARTSEITE:04d}.png")
    print("\n   Wenn der Screenshot nicht gut ist:")
    print("   1. Passe CLIP-Koordinaten in Zeile 85-90 an")
    print("   2. Oder nutze 'python finde_koordinaten.py'")
    print("   3. Teste mit ANZAHL_SEITEN = 1 bis perfekt")




if __name__ == "__main__":
    hauptprogramm()